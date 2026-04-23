#!/usr/bin/env python3
"""Simulated IRB Reviewer — reads generated forms and produces review opinions.

Applies 45 CFR 46.111 criteria and institutional review patterns to
catch errors, inconsistencies, and missing elements before real submission.
"""
import os
import sys
import glob
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from scripts.docx_utils import load_config
from scripts.form_selector import select_forms, FORM_REGISTRY
from scripts.review_criteria import (
    REVIEW_CRITERIA, PLACEHOLDER_PATTERNS, DECISIONS, SEVERITY,
)
from scripts.institution_profiles import get_institution_profile, get_phase_name


def extract_text(docx_path):
    """Extract all text from a DOCX file (paragraphs + table cells)."""
    doc = Document(docx_path)
    texts = []
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text)
    return "\n".join(texts)


def find_placeholders(text, form_id):
    """Find placeholder strings that indicate incomplete content."""
    findings = []
    for pattern in PLACEHOLDER_PATTERNS:
        if pattern in text:
            count = text.count(pattern)
            findings.append(
                f"**{form_id}**: Found placeholder '{pattern}' ({count} occurrence{'s' if count > 1 else ''})"
            )
    return findings


def check_completeness(config, output_dir, form_texts):
    """Check document completeness criteria."""
    results = []
    findings = []

    # Check all required forms are generated
    expected = select_forms(config)
    expected_ids = [fid for fid, _ in expected]
    generated_ids = list(form_texts.keys())
    missing = [fid for fid in expected_ids if fid not in generated_ids]

    if missing:
        results.append(("forms_complete", False, f"Missing forms: {', '.join(missing)}"))
        findings.append(("required", f"Missing required forms: {', '.join(missing)}"))
    else:
        results.append(("forms_complete", True, f"All required forms generated ({len(generated_ids)}/{len(expected_ids)})"))

    # Check IRB number consistency
    irb_no = config["study"]["irb_no"]
    irb_missing = []
    for fid, text in form_texts.items():
        if irb_no and irb_no not in text:
            irb_missing.append(fid)
    if irb_missing:
        results.append(("irb_no_consistent", False, f"IRB number '{irb_no}' not found in: {', '.join(irb_missing)}"))
        findings.append(("required", f"IRB number '{irb_no}' missing from: {', '.join(irb_missing)}"))
    else:
        results.append(("irb_no_consistent", True, f"IRB number '{irb_no}' found in {len(form_texts)}/{len(form_texts)} forms"))

    # Check title consistency
    title_zh = config["study"]["title_zh"]
    title_missing = []
    if title_zh:
        # Check first 20 chars to handle line breaks
        title_fragment = title_zh[:20]
        for fid, text in form_texts.items():
            if title_fragment not in text:
                title_missing.append(fid)
    if title_missing:
        results.append(("title_consistent", False, f"Title fragment not found in: {', '.join(title_missing)}"))
        findings.append(("suggestion", f"Study title not found in: {', '.join(title_missing)} — may be truncated or formatted differently"))
    else:
        results.append(("title_consistent", True, "Title consistent across all forms"))

    # Check PI name
    pi_name = config["pi"]["name"]
    pi_missing = []
    for fid, text in form_texts.items():
        if pi_name and pi_name not in text:
            pi_missing.append(fid)
    if pi_missing:
        results.append(("pi_consistent", False, f"PI name '{pi_name}' not found in: {', '.join(pi_missing)}"))
        findings.append(("required", f"PI name '{pi_name}' missing from: {', '.join(pi_missing)}"))
    else:
        results.append(("pi_consistent", True, f"PI name '{pi_name}' found in all forms"))

    # Check dates
    dates_blank = not config["dates"].get("study_start") or not config["dates"].get("study_end")
    if dates_blank:
        results.append(("dates_present", False, "Study start or end date is blank"))
        findings.append(("required", "Study dates (start/end) must be specified in config.yml"))
    else:
        results.append(("dates_present", True, f"Study period: {config['dates']['study_start']} — {config['dates']['study_end']}"))

    # Check co-PI
    co_pis = config.get("co_pi", [])
    if co_pis:
        for cp in co_pis:
            cp_missing = []
            for fid, text in form_texts.items():
                if fid in ["SF001", "SF002", "SF037"] and cp["name"] not in text:
                    cp_missing.append(fid)
            if cp_missing:
                results.append(("copi_listed", False, f"Co-PI {cp['name']} not found in: {', '.join(cp_missing)}"))
                findings.append(("suggestion", f"Co-PI {cp['name']} should appear in {', '.join(cp_missing)}"))
            else:
                results.append(("copi_listed", True, f"Co-PI {cp['name']} listed in relevant forms"))
    else:
        results.append(("copi_listed", True, "No co-PI — N/A"))

    # Check footers
    results.append(("footer_present", True, "Form version footers present (auto-generated)"))

    return results, findings


def check_consent(config, form_texts):
    """Check informed consent criteria."""
    results = []
    findings = []
    is_retro = config["study"].get("type") == "retrospective"
    has_waiver = config["subjects"].get("consent_waiver", False)

    if is_retro and has_waiver:
        if "SF005" in form_texts:
            results.append(("consent_waiver_justified", True, "Consent waiver form (SF005) included — appropriate for retrospective study"))
        else:
            results.append(("consent_waiver_justified", False, "Retrospective study with consent waiver but SF005 not generated"))
            findings.append(("required", "Consent waiver form (SF005) required for retrospective studies with consent_waiver=true"))
    elif not has_waiver:
        consent_forms = [fid for fid in form_texts if fid in ["SF062", "SF063", "SF075", "SF090"]]
        if consent_forms:
            results.append(("consent_elements", True, f"Consent form(s) included: {', '.join(consent_forms)}"))
            # Check for contact info in consent
            for fid in consent_forms:
                text = form_texts[fid]
                if config["pi"]["phone"] not in text and config["pi"]["email"] not in text:
                    findings.append(("suggestion", f"**{fid}**: PI contact info (phone/email) should appear in consent form"))
        else:
            results.append(("consent_elements", False, "No consent form generated but consent_waiver is false"))
            findings.append(("required", "Consent form required when consent_waiver is false"))

    # Signature blocks
    results.append(("signature_blocks", True, "Signature blocks present (auto-generated)"))
    results.append(("contact_info", True, "Contact info included where applicable"))

    return results, findings


def check_risk_benefit(config, form_texts):
    """Check risk-benefit assessment criteria."""
    results = []
    findings = []
    is_retro = config["study"].get("type") == "retrospective"
    review_type = config["study"].get("review_type", "")

    if is_retro and review_type == "expedited":
        results.append(("risk_level_match", True, "Minimal risk — expedited review appropriate for retrospective study"))
    elif review_type == "full_board":
        results.append(("risk_level_match", True, "Full board review — matches greater than minimal risk"))
    elif review_type == "exempt":
        results.append(("risk_level_match", True, "Exempt review — minimal risk research"))
    else:
        results.append(("risk_level_match", False, f"Review type '{review_type}' — verify risk level justification"))
        findings.append(("suggestion", f"Verify that review type '{review_type}' matches actual risk level"))

    # Check for risk description in SF002
    if "SF002" in form_texts:
        text = form_texts["SF002"]
        if is_retro:
            results.append(("risk_minimization", True, "Retrospective study — minimal risk inherent in design"))
        else:
            results.append(("risk_minimization", False, "Verify risk minimization strategies are described in protocol"))
            findings.append(("suggestion", "Ensure risk minimization strategies are detailed in study protocol"))

    results.append(("benefits_stated", True, "Benefits addressed in study design"))

    return results, findings


def check_privacy(config, form_texts):
    """Check privacy and confidentiality criteria."""
    results = []
    findings = []
    closure_cfg = config.get("closure", {})
    data_safety = closure_cfg.get("data_safety", {})

    if data_safety.get("deidentified"):
        results.append(("deidentification", True, "Data de-identification confirmed in config"))
    else:
        results.append(("deidentification", False, "De-identification not confirmed"))
        findings.append(("required", "Specify de-identification method for subject data"))

    if data_safety.get("encrypted"):
        results.append(("data_security", True, "Data encryption confirmed"))
    else:
        results.append(("data_security", False, "Data encryption not confirmed"))
        findings.append(("suggestion", "Consider encrypting research data at rest and in transit"))

    retention = data_safety.get("retention_years")
    if retention:
        results.append(("retention_period", True, f"Data retention: {retention} years"))
    else:
        results.append(("retention_period", False, "Data retention period not specified"))
        findings.append(("required", "Specify data retention period in config.yml closure.data_safety.retention_years"))

    personnel = data_safety.get("authorized_personnel")
    if personnel:
        results.append(("access_control", True, f"Authorized personnel: {personnel}"))
    else:
        results.append(("access_control", False, "Authorized personnel not specified"))
        findings.append(("suggestion", "Specify who has access to research data"))

    return results, findings


def check_subject_selection(config, form_texts):
    """Check subject selection criteria."""
    results = []
    findings = []

    planned_n = config["subjects"].get("planned_n", 0)
    if planned_n > 0:
        results.append(("sample_size", True, f"Sample size: {planned_n}"))
    else:
        results.append(("sample_size", False, "Sample size is 0 or not specified"))
        findings.append(("required", "Specify planned sample size (subjects.planned_n)"))

    vuln = config["subjects"].get("vulnerable_population", False)
    results.append(("vulnerable_flag", True, f"Vulnerable population: {'Yes — additional protections needed' if vuln else 'No'}"))
    if vuln:
        findings.append(("suggestion", "Ensure additional safeguards for vulnerable populations are documented in protocol"))

    results.append(("selection_equitable", True, "Subject selection addressed in study design"))

    return results, findings


def check_study_design(config, form_texts):
    """Check study design criteria."""
    results = []
    findings = []

    study_type = config["study"].get("type", "")
    phase = config.get("phase", "")

    # Check type-form alignment
    if study_type == "retrospective":
        if "SF005" in form_texts or "SF003" in form_texts:
            results.append(("type_forms_match", True, "Retrospective study — waiver/expedited forms correctly included"))
        else:
            results.append(("type_forms_match", False, "Retrospective study but no waiver/expedited forms"))
    else:
        results.append(("type_forms_match", True, f"Study type '{study_type}' — form set verified"))

    # Data period
    data_period = config["dates"].get("data_period", "")
    if study_type == "retrospective" and not data_period:
        results.append(("data_period_scope", False, "Retrospective study but no data_period specified"))
        findings.append(("required", "Specify data collection period for retrospective study (dates.data_period)"))
    else:
        results.append(("data_period_scope", True, f"Data period: {data_period}" if data_period else "N/A for prospective study"))

    results.append(("design_consistent", True, f"Study design: {config['study'].get('design', 'N/A')}"))

    return results, findings


def check_administrative(config, form_texts, institution_profile):
    """Check administrative compliance criteria."""
    results = []
    findings = []

    if "SF094" in form_texts:
        results.append(("financial_disclosure", True, "Financial disclosure (SF094) included"))
        # Check if co-PIs need separate disclosures
        co_pis = config.get("co_pi", [])
        for cp in co_pis:
            if cp["name"] not in form_texts.get("SF094", ""):
                findings.append(("suggestion", f"Co-PI {cp['name']} may need a separate SF094 financial disclosure"))
    else:
        phase = config.get("phase", "")
        if phase in ["new", "amendment", "continuing"]:
            results.append(("financial_disclosure", False, "SF094 not generated but required for this phase"))
            findings.append(("required", "Financial disclosure form (SF094) required"))
        else:
            results.append(("financial_disclosure", True, "SF094 not required for this phase"))

    results.append(("all_signatures", True, "Signature blocks present (manual step required)"))
    results.append(("current_versions", True, "Using current form versions"))
    results.append(("submission_noted", True, f"Submit to {institution_profile['submission_email']}"))

    return results, findings


def check_rules_of_thumb(config, form_texts):
    """Apply experienced reviewer rules of thumb."""
    findings = []
    is_retro = config["study"].get("type") == "retrospective"
    review_type = config["study"].get("review_type", "")
    drug_device = config["study"].get("drug_device", False)
    phase = config.get("phase", "")

    # Rule: clinical trial claiming expedited is a red flag
    if config["study"].get("type") == "clinical_trial" and review_type == "expedited":
        findings.append(("required",
            "Clinical trial marked as expedited review — clinical trials with novel interventions "
            "typically require full board review (45 CFR 46.111)"))

    # Rule: drug/device study without consent forms
    if drug_device and config["subjects"].get("consent_waiver"):
        findings.append(("required",
            "Drug/device study with consent waiver — interventional studies "
            "require informed consent (consent waiver typically only for retrospective/minimal risk)"))

    # Rule: large sample with no groups defined
    planned_n = config["subjects"].get("planned_n", 0)
    groups = config["subjects"].get("groups", [])
    if planned_n > 100 and not groups:
        findings.append(("suggestion",
            f"Sample size is {planned_n} but no study groups defined — "
            "consider defining comparison groups for statistical analysis"))

    # Rule: retrospective without data period
    if is_retro and not config["dates"].get("data_period"):
        findings.append(("required",
            "Retrospective study without data collection period — "
            "specify the date range of medical records being reviewed"))

    # Rule: consent waiver without all 4 criteria justification
    if config["subjects"].get("consent_waiver") and phase == "new":
        if "SF005" not in form_texts:
            findings.append(("required",
                "Consent waiver claimed but SF005 (免取得知情同意檢核表) not generated — "
                "all 4 waiver criteria must be documented (45 CFR 46.116(f)(3))"))

    # Rule: title too vague or missing
    title = config["study"].get("title_zh", "")
    if len(title) < 10:
        findings.append(("required",
            "Study title is too short or missing — provide a descriptive Chinese title"))

    # Rule: study period > 3 years for retrospective (unusual)
    start = config["dates"].get("study_start", "")
    end = config["dates"].get("study_end", "")
    if start and end and is_retro:
        try:
            start_y = int(start[:4]) if start[:4].isdigit() else 0
            end_y = int(end[:4]) if end[:4].isdigit() else 0
            if end_y - start_y > 3:
                findings.append(("note",
                    f"Study period spans {end_y - start_y} years — "
                    "longer study periods may require continuing review; plan for annual renewal"))
        except (ValueError, IndexError):
            pass

    # Rule: co-PI from different department (strength, note it)
    co_pis = config.get("co_pi", [])
    pi_dept = config["pi"].get("dept", "")
    for cp in co_pis:
        cp_dept = cp.get("dept", "")
        if cp_dept and pi_dept and cp_dept.split("／")[0] != pi_dept.split("／")[0]:
            findings.append(("note",
                f"Cross-departmental collaboration: PI ({pi_dept.split('／')[0]}) + "
                f"Co-PI {cp['name']} ({cp_dept.split('／')[0]}) — "
                "strengthens multidisciplinary perspective"))

    # Rule: no IRB number for non-new submissions
    if phase != "new" and not config["study"].get("irb_no"):
        findings.append(("required",
            f"Phase is '{phase}' but IRB number is blank — "
            "non-new submissions require an existing IRB approval number"))

    return findings


def run_review(config_path="config.yml", output_dir="output"):
    """Run full review and generate opinion markdown."""
    config = load_config(config_path)
    irb_no = config["study"]["irb_no"]
    phase = config["phase"]
    institution_profile = get_institution_profile(config)
    phase_zh = get_phase_name(phase, institution_profile)

    # Load all generated DOCX files (recursive for harness mode)
    docx_files = sorted(glob.glob(os.path.join(output_dir, "**", "*.docx"), recursive=True))
    if not docx_files:
        print("✗ No DOCX files found in output/. Run `make generate` first.")
        sys.exit(1)

    form_texts = {}
    for f in docx_files:
        basename = os.path.basename(f)
        # Extract form ID (e.g., "SF001" from "SF001_20260401A.docx")
        fid = basename.split("_")[0]
        form_texts[fid] = extract_text(f)

    print(f"Reviewing {len(form_texts)} forms for {phase_zh} ({irb_no})...")

    # Run all checks
    all_results = {}
    all_findings = []

    checkers = [
        ("completeness", check_completeness, (config, output_dir, form_texts)),
        ("consent", check_consent, (config, form_texts)),
        ("risk_benefit", check_risk_benefit, (config, form_texts)),
        ("privacy", check_privacy, (config, form_texts)),
        ("subject_selection", check_subject_selection, (config, form_texts)),
        ("study_design", check_study_design, (config, form_texts)),
        ("administrative", check_administrative, (config, form_texts, institution_profile)),
    ]

    for category, checker, args in checkers:
        results, findings = checker(*args)
        all_results[category] = results
        all_findings.extend(findings)

    # Run rules of thumb
    rot_findings = check_rules_of_thumb(config, form_texts)
    all_findings.extend(rot_findings)

    # Find placeholders in all forms
    placeholder_findings = []
    for fid, text in form_texts.items():
        placeholder_findings.extend(find_placeholders(text, fid))
    if placeholder_findings:
        for pf in placeholder_findings:
            all_findings.append(("suggestion", pf))

    # Determine decision
    required_count = sum(1 for sev, _ in all_findings if sev == "required")
    if required_count == 0:
        decision = DECISIONS["approved"]
    elif required_count <= 3:
        decision = DECISIONS["approved_with_revisions"]
    else:
        decision = DECISIONS["tabled"]

    # Generate report
    now = datetime.now()
    timestamp = now.strftime("%Y%m%d_%H%M%S")
    date_str = now.strftime("%Y-%m-%d")

    form_list = ", ".join(sorted(form_texts.keys()))
    title_display = config["study"]["title_zh"][:40] + "..." if len(config["study"]["title_zh"]) > 40 else config["study"]["title_zh"]

    lines = [
        "# IRB 審查意見書",
        "## 模擬審查 — IRB-in-Hurry Reviewer",
        "",
        "| 項目 | 內容 |",
        "|------|------|",
        f"| IRB 編號 | {irb_no} |",
        f"| 計畫名稱 | {title_display} |",
        f"| 主持人 | {config['pi']['name']} |",
        f"| 審查階段 | {phase_zh} |",
        f"| 審查日期 | {date_str} |",
        f"| 產生表單 | {form_list} |",
        "",
        "---",
        "",
        "## 審查決定",
        "",
        f"**{decision}**",
        "",
        "---",
        "",
        "## 審查清單",
        "",
    ]

    # Criteria results
    criteria_num = {
        "completeness": "一", "consent": "二", "risk_benefit": "三",
        "privacy": "四", "subject_selection": "五", "study_design": "六",
        "administrative": "七",
    }
    for category, criteria in REVIEW_CRITERIA.items():
        num = criteria_num.get(category, "")
        lines.append(f"### {num}、{criteria['name']} ({criteria['name_en']})")
        lines.append(f"*{criteria['regulatory_basis']}*")
        lines.append("")
        if category in all_results:
            for item_id, passed, detail in all_results[category]:
                mark = "■" if passed else "□"
                lines.append(f"- {mark} {detail}")
        lines.append("")

    # Findings by severity
    lines.extend(["---", ""])

    for severity_key, severity_label in SEVERITY.items():
        items = [(sev, msg) for sev, msg in all_findings if sev == severity_key]
        if items:
            lines.append(f"## {severity_label}")
            lines.append("")
            for i, (_, msg) in enumerate(items, 1):
                lines.append(f"{i}. {msg}")
            lines.append("")

    if not all_findings:
        lines.extend([
            "## 審查結果",
            "",
            "未發現需要修正的事項。所有表單符合送審要求。",
            "",
        ])

    # 45 CFR 46.111 summary
    is_retro = config["study"].get("type") == "retrospective"
    review_type = config["study"].get("review_type", "")
    lines.extend([
        "---",
        "",
        "## 45 CFR 46.111 核准要件摘要",
        "",
        f"| # | Criterion | Status |",
        f"|---|-----------|--------|",
        f"| 1 | Risk minimization | {'■ Retrospective design — minimal risk' if is_retro else '□ Verify in protocol'} |",
        f"| 2 | Reasonable risk-benefit | {'■ Knowledge benefit, no subject risk' if is_retro else '□ Verify in protocol'} |",
        f"| 3 | Equitable selection | {'■ Consecutive patients from chart review' if is_retro else '□ Verify recruitment plan'} |",
        f"| 4 | Informed consent sought | {'■ Waiver justified (retrospective)' if config['subjects'].get('consent_waiver') else '□ Consent form required'} |",
        f"| 5 | Consent documented | {'■ SF005 waiver checklist' if config['subjects'].get('consent_waiver') else '□ Signed consent'} |",
        f"| 6 | Safety monitoring | {'■ PI monitors data quality' if is_retro else '□ DSMP required'} |",
        f"| 7 | Privacy protected | {'■' if config.get('closure', {}).get('data_safety', {}).get('deidentified') else '□'} De-identification + encryption |",
        f"| 8 | Vulnerable safeguards | {'■ N/A' if not config['subjects'].get('vulnerable_population') else '□ Additional protections needed'} |",
        "",
    ])

    # Footer
    lines.extend([
        "---",
        "",
        f"*本審查意見由 IRB-in-Hurry Reviewer 自動產生，僅供參考。*",
        f"*{institution_profile['decision_note']}*",
        f"*送審請寄：{institution_profile['submission_email']}*",
        f"*審查指引：see .claude/skills/irb/references/reviewer-guide.md*",
    ])

    # Write to reviewers/
    review_dir = "reviewers"
    os.makedirs(review_dir, exist_ok=True)
    review_path = os.path.join(review_dir, f"review_{timestamp}.md")
    with open(review_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")

    # Print summary
    suggestion_count = sum(1 for sev, _ in all_findings if sev == "suggestion")
    note_count = sum(1 for sev, _ in all_findings if sev == "note")

    print(f"\n{'═' * 50}")
    print(f"  Decision: {decision}")
    print(f"  Required:    {required_count}")
    print(f"  Suggestions: {suggestion_count}")
    print(f"  Notes:       {note_count}")
    print(f"  Report:      {review_path}")
    print(f"{'═' * 50}")

    return review_path


if __name__ == "__main__":
    config_path = sys.argv[1] if len(sys.argv) > 1 else "config.yml"
    run_review(config_path)
