"""Generate 中文計畫摘要 (Chinese Project Summary) for KFSYSCC IRB.

Required for new case submission. Max 2 pages.
11 sections per KFSYSCC official template.
"""
import os
from scripts.docx_utils import (
    init_doc, add_p, add_header, add_footer,
    set_run_font, add_ct, apply_tb, set_cell_shading, check,
)
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _study_type_zh(config):
    """Get Chinese description of study type."""
    t = config["study"].get("type", "")
    d = config["study"].get("design", "")
    mapping = {
        "retrospective": "回溯性",
        "prospective": "前瞻性",
        "clinical_trial": "臨床試驗",
        "genetic": "基因研究",
    }
    design_mapping = {
        "cohort": "世代研究",
        "case_control": "病例對照研究",
        "cross_sectional": "橫斷面研究",
        "rct": "隨機對照試驗",
        "single_arm": "單臂研究",
    }
    type_zh = mapping.get(t, t)
    design_zh = design_mapping.get(d, d)
    return f"{type_zh}{design_zh}"


def generate_proposal_summary(config, output_dir):
    """Generate 中文計畫摘要 (Chinese Project Summary).

    11-section template per KFSYSCC requirements. Max 2 pages.
    """
    doc = init_doc(11)

    # Compact margins for 2-page limit
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", True, 13,
          WD_ALIGN_PARAGRAPH.CENTER, Pt(2))
    add_p(doc, "中文計畫摘要", True, 15,
          WD_ALIGN_PARAGRAPH.CENTER, Pt(8))

    # Header info table
    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl.rows[0].cells[0], "KFSYSCC-IRB編號", True, 10)
    add_ct(tbl.rows[0].cells[1], config["study"]["irb_no"] or "（待核發）", size=10)
    add_ct(tbl.rows[1].cells[0], "計畫名稱", True, 10)
    add_ct(tbl.rows[1].cells[1], config["study"]["title_zh"], size=10)
    add_ct(tbl.rows[2].cells[0], "計畫主持人", True, 10)
    pi_str = config["pi"]["name"]
    co_pis = config.get("co_pi", [])
    if co_pis:
        pi_str += "（主持人）、" + "、".join(
            f"{cp['name']}（共同主持人）" for cp in co_pis
        )
    add_ct(tbl.rows[2].cells[1], pi_str, size=10)
    apply_tb(tbl)

    doc.add_paragraph()

    is_retro = config["study"].get("type") == "retrospective"
    study_type_zh = _study_type_zh(config)
    data_period = config["dates"].get("data_period", "")
    planned_n = config["subjects"].get("planned_n", 0)

    # Section 一、研究主題
    add_p(doc, "一、研究主題", True, 11, sa=Pt(4), sb=Pt(4))
    add_p(doc, config["study"]["title_zh"], size=10, sa=Pt(2))

    # Section 二、研究背景
    add_p(doc, "二、研究背景", True, 11, sa=Pt(4), sb=Pt(4))
    add_p(doc, "（請簡述疾病現況、自然病程、現有治療方式及預後，說明本研究之必要性。）",
          size=10, sa=Pt(2))

    # Section 三、研究目的
    add_p(doc, "三、研究目的", True, 11, sa=Pt(4), sb=Pt(4))
    add_p(doc, "（請說明主要研究目的及次要研究目的。）", size=10, sa=Pt(2))

    # Section 四、執行期間
    add_p(doc, "四、執行期間", True, 11, sa=Pt(4), sb=Pt(4))
    period_text = f"計畫期間：{config['dates'].get('study_start', '')} 至 {config['dates'].get('study_end', '')}"
    if is_retro and data_period:
        period_text += f"\n資料回溯期間：{data_period}"
    add_p(doc, period_text, size=10, sa=Pt(2))

    # Section 五、研究設計
    add_p(doc, "五、研究設計", True, 11, sa=Pt(4), sb=Pt(4))
    design_text = f"本研究採{study_type_zh}設計"
    if is_retro:
        design_text += f"，回溯分析{data_period}期間之病歷資料。"
        design_text += "\n本研究為非介入性研究，不涉及任何實驗性處置。"
    else:
        design_text += "。\n（請說明對照組、盲性、隨機分組等設計。）"
    add_p(doc, design_text, size=10, sa=Pt(2))

    # Section 六、研究參與者
    add_p(doc, "六、研究參與者", True, 11, sa=Pt(4), sb=Pt(4))
    subj_text = f"預計收錄人數：{planned_n}人"
    groups = config["subjects"].get("groups", [])
    if groups:
        subj_text += "\n分組："
        for g in groups:
            subj_text += f"\n  - {g['name']}"
            if g.get("n"):
                subj_text += f"：{g['n']}人"
    subj_text += "\n\n納入條件：（請列出）"
    subj_text += "\n排除條件：（請列出）"
    if config["subjects"].get("consent_waiver"):
        subj_text += "\n\n受試者保護措施：本研究為回溯性病歷審查，經IRB核准免取得知情同意。所有資料均去識別化處理。"
    else:
        subj_text += "\n\n受試者保護措施：（請說明知情同意取得方式及受試者保護機制。）"
    add_p(doc, subj_text, size=10, sa=Pt(2))

    # Section 七、研究方法
    add_p(doc, "七、研究方法", True, 11, sa=Pt(4), sb=Pt(4))
    if is_retro:
        add_p(doc, ("資料來源：和信治癌中心醫院電子病歷系統\n"
                    "資料收集項目：（請列出收集之臨床變項）\n"
                    "統計方法：（請說明使用之統計分析方法）"),
              size=10, sa=Pt(2))
    else:
        add_p(doc, "（請說明治療程序、劑量、臨床觀察、追蹤時程及療效評估。）",
              size=10, sa=Pt(2))

    # Section 八、不良事件處理
    add_p(doc, "八、不良事件處理", True, 11, sa=Pt(4), sb=Pt(4))
    if is_retro:
        add_p(doc, "不適用。本研究為回溯性病歷審查研究，非介入性研究，無不良事件通報需求。",
              size=10, sa=Pt(2))
    else:
        add_p(doc, "（請說明不良事件的發生率及處理原則。）", size=10, sa=Pt(2))

    # Section 九、統計分析
    add_p(doc, "九、統計分析", True, 11, sa=Pt(4), sb=Pt(4))
    add_p(doc, "（請說明統計分析計畫，包含期中分析（如適用）。）",
          size=10, sa=Pt(2))

    # Section 十、資料保護及安全監測
    add_p(doc, "十、資料保護及安全監測", True, 11, sa=Pt(4), sb=Pt(4))
    ds = config.get("closure", {}).get("data_safety", {})
    safety_text = ""
    if ds.get("deidentified"):
        safety_text += "■ 所有資料均去識別化處理\n"
    if ds.get("encrypted"):
        safety_text += "■ 電子資料以加密方式儲存，設有密碼保護\n"
    if ds.get("retention_years"):
        safety_text += f"■ 資料保存期限：研究結束後{ds['retention_years']}年\n"
    if ds.get("authorized_personnel"):
        safety_text += f"■ 資料存取授權人員：{ds['authorized_personnel']}\n"
    if not safety_text:
        safety_text = "（請說明資料保密措施及受試者安全監測計畫。）"
    add_p(doc, safety_text.strip(), size=10, sa=Pt(2))

    # Section 十一、附件
    add_p(doc, "十一、附件", True, 11, sa=Pt(4), sb=Pt(4))
    add_p(doc, "（如有相關附件，請列出。）", size=10, sa=Pt(2))

    # Footer note
    doc.add_paragraph()
    add_p(doc, "※ 本摘要以不超過2頁為原則。", size=9,
          alignment=WD_ALIGN_PARAGRAPH.RIGHT, sa=Pt(2))

    irb_no = config["study"]["irb_no"] or "proposal"
    path = os.path.join(output_dir, f"中文計畫摘要_{irb_no}.docx")
    doc.save(path)
    return path


ALL_GENERATORS = {
    "PROPOSAL": generate_proposal_summary,
}
