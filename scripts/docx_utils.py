"""Shared DOCX helper functions for IRB form generation.

Extracted from irb-close/generate_forms.py, refactored to accept config dict.
"""
import os
import yaml
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def load_config(path="config.yml"):
    """Load and return config dict from YAML file."""
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def check(condition: bool) -> str:
    """Return ■ if True, □ if False (IRB checkbox convention)."""
    return "■" if condition else "□"


def init_doc(sz=12):
    """Create a new Document with 標楷體 default font."""
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = '標楷體'
    s.font.size = Pt(sz)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    return doc


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, start, end with {val, sz, color}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", 4)}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, font_name="標楷體", size=12, bold=False):
    """Set run font with eastAsia fallback."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run.font.element.rPr
    if rPr is None:
        run.font.element.get_or_add_rPr()
        rPr = run.font.element.rPr
    rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_p(doc, text, bold=False, size=12, alignment=None, sa=Pt(6), sb=Pt(0)):
    """Add a formatted paragraph to the document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "標楷體", size, bold)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = sa
    p.paragraph_format.space_before = sb
    return p


def add_ct(cell, text, bold=False, size=10, alignment=None):
    """Add formatted text to a table cell."""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear()
    run = p.add_run(text)
    set_run_font(run, "標楷體", size, bold)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    return p


def apply_tb(tbl):
    """Apply borders to all cells in a table."""
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 4}, bottom={"sz": 4},
                start={"sz": 4}, end={"sz": 4},
            )


def add_header(doc, config, inc_proj=True):
    """Add standard IRB header block (IRB no, title, PI, dept)."""
    rows = 5 if inc_proj else 4
    tbl = doc.add_table(rows=rows, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    i = 0
    add_ct(tbl.rows[i].cells[0], "KFSYSCC-IRB編號", True, 11)
    add_ct(tbl.rows[i].cells[1], config["study"]["irb_no"], size=11)

    if inc_proj:
        i += 1
        add_ct(tbl.rows[i].cells[0], "計畫編號", True, 11)
        add_ct(tbl.rows[i].cells[1], config["study"].get("project_no", "不適用"), size=11)

    i += 1
    add_ct(tbl.rows[i].cells[0], "計畫名稱", True, 11)
    cell = tbl.rows[i].cells[1]
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(f"（中文）{config['study']['title_zh']}")
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(2)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(f"（英文）{config['study']['title_en']}")
    set_run_font(run2, size=10)
    p2.paragraph_format.space_after = Pt(2)

    i += 1
    # Build PI/co-PI string
    pi_str = config["pi"]["name"]
    co_pis = config.get("co_pi", [])
    if co_pis:
        pi_str += "（主持人）"
        for cp in co_pis:
            pi_str += f"\n{cp['name']}（共同主持人）"
    add_ct(tbl.rows[i].cells[0], "計畫主持人", True, 11)
    add_ct(tbl.rows[i].cells[1], pi_str, size=11)

    i += 1
    add_ct(tbl.rows[i].cells[0], "單位／職稱", True, 11)
    add_ct(tbl.rows[i].cells[1], config["pi"]["dept"], size=11)

    apply_tb(tbl)
    doc.add_paragraph()


def add_footer(doc, ver, fno, date):
    """Add version/form number footer."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"版次第{ver}版　{fno}　{date}")
    set_run_font(run, size=9)


def _replace_text_in_paragraph(paragraph, replacements):
    """Replace known text snippets in a paragraph.

    Paragraph-level replacement is pragmatic: it may flatten complex run styling for
    replaced content, but it keeps the existing generator output shape stable.
    """
    if not replacements:
        return
    original = paragraph.text
    updated = original
    for old_text, new_text in replacements.items():
        updated = updated.replace(old_text, new_text)
    if updated == original:
        return

    if not paragraph.runs:
        paragraph.add_run(updated)
        return

    first = paragraph.runs[0]
    first.text = updated
    for run in paragraph.runs[1:]:
        run.text = ""


def apply_institution_text_replacements(docx_path, institution_profile):
    """Apply institution text replacements to generated DOCX in-place."""
    from docx import Document

    replacements = institution_profile.get("text_replacements", {})
    if not replacements:
        return

    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, replacements)

    doc.save(docx_path)
