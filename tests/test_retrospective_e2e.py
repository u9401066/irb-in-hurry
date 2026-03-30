"""End-to-end test: retrospective study → generate all forms → validate output."""
import pytest
import yaml
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from scripts.form_selector import select_forms
from scripts.docx_utils import load_config


@pytest.fixture
def retro_config():
    return load_config("tests/fixtures/sample_retrospective.yml")


@pytest.fixture
def output_dir(tmp_path):
    return str(tmp_path / "output")


def test_new_case_generates_all_forms(retro_config, output_dir):
    """Generate new case forms and verify all files created."""
    import importlib
    from scripts.form_selector import get_generator

    retro_config["phase"] = "new"
    os.makedirs(output_dir, exist_ok=True)

    forms = select_forms(retro_config)
    generated = []
    for fid, name_zh in forms:
        gen_info = get_generator(fid)
        assert gen_info is not None, f"No generator for {fid}"
        mod_path, func_name = gen_info
        mod = importlib.import_module(f"scripts.{mod_path}")
        gen_func = getattr(mod, func_name)
        path = gen_func(retro_config, output_dir)
        assert os.path.exists(path), f"{fid} file not created: {path}"
        assert os.path.getsize(path) > 0, f"{fid} file is empty: {path}"
        generated.append(path)

    assert len(generated) == 6  # SF001, SF002, SF094, PROPOSAL, SF003, SF005


def test_closure_generates_all_forms(retro_config, output_dir):
    """Generate closure forms and verify all files created."""
    import importlib
    from scripts.form_selector import get_generator

    retro_config["phase"] = "closure"
    os.makedirs(output_dir, exist_ok=True)

    forms = select_forms(retro_config)
    generated = []
    for fid, name_zh in forms:
        gen_info = get_generator(fid)
        assert gen_info is not None, f"No generator for {fid}"
        mod_path, func_name = gen_info
        mod = importlib.import_module(f"scripts.{mod_path}")
        gen_func = getattr(mod, func_name)
        path = gen_func(retro_config, output_dir)
        assert os.path.exists(path), f"{fid} file not created: {path}"
        assert os.path.getsize(path) > 0, f"{fid} file is empty: {path}"
        generated.append(path)

    assert len(generated) == 4  # SF036, SF037, SF038, SF023


def test_docx_contains_irb_number(retro_config, output_dir):
    """Verify generated DOCX files contain the IRB number."""
    import importlib
    from scripts.form_selector import get_generator
    from docx import Document

    retro_config["phase"] = "new"
    os.makedirs(output_dir, exist_ok=True)

    gen_info = get_generator("SF001")
    mod = importlib.import_module(f"scripts.{gen_info[0]}")
    path = getattr(mod, gen_info[1])(retro_config, output_dir)

    doc = Document(path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    assert "20250801A" in all_text, "IRB number not found in SF001"


def test_docx_contains_pi_name(retro_config, output_dir):
    """Verify generated DOCX contains PI name."""
    import importlib
    from scripts.form_selector import get_generator
    from docx import Document

    retro_config["phase"] = "new"
    os.makedirs(output_dir, exist_ok=True)

    gen_info = get_generator("SF002")
    mod = importlib.import_module(f"scripts.{gen_info[0]}")
    path = getattr(mod, gen_info[1])(retro_config, output_dir)

    doc = Document(path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    assert "林協霆" in all_text, "PI name not found in SF002"


def test_checklist_generation(retro_config, output_dir, tmp_path):
    """Verify checklist.md is generated correctly."""
    from scripts.checklist import generate_checklist

    results = [
        ("SF001", "新案審查送審資料表", "/fake/path.docx", "generated"),
        ("SF002", "研究計畫申請書", "/fake/path2.docx", "generated"),
        ("SF094", "顯著財務利益申報表", None, "error"),
    ]
    checklist_path = str(tmp_path / "checklist.md")
    generate_checklist(retro_config, results, "新案審查", checklist_path)

    with open(checklist_path) as f:
        content = f.read()

    assert "20250801A" in content
    assert "■ SF001" in content  # generated
    assert "■ SF002" in content  # generated
    assert "□ SF094" in content  # error
    assert "irb@kfsyscc.org" in content


def test_config_validation():
    """Verify config loads without error."""
    config = load_config("tests/fixtures/sample_retrospective.yml")
    assert config["study"]["irb_no"] == "20250801A"
    assert config["pi"]["name"] == "林協霆"
    assert config["subjects"]["consent_waiver"] is True
    assert len(config["co_pi"]) == 1
    assert config["co_pi"][0]["name"] == "邱倫維"
